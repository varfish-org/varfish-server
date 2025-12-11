"""API views for exporting filter settings."""

import datetime
from io import BytesIO
import json
import logging
import re

from django.contrib.auth.models import AnonymousUser
from django.forms import model_to_dict
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from projectroles.models import Project, RoleAssignment
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from variants.models import Case
from variants.models.presets import (
    ChromosomePresets,
    FlagsEtcPresets,
    FrequencyPresets,
    ImpactPresets,
    PresetSet,
    QualityPresets,
    QuickPresets,
)

logger = logging.getLogger(__name__)


def _check_user_authentication(request):
    """
    Check if the user is authenticated.

    Args:
        request: Django request object

    Returns:
        JsonResponse or None: Error response if not authenticated, None if authenticated
    """
    if isinstance(request.user, AnonymousUser) or not request.user.is_authenticated:
        return JsonResponse({"error": "Authentication required"}, status=401)
    return None


def _check_project_permission_by_uuid(request, project_uuid):
    """
    Check if the user has permission to access a project by UUID.

    Args:
        request: Django request object
        project_uuid: UUID string of the project to check access for

    Returns:
        tuple: (has_permission: bool, project: Project or None, error_response: JsonResponse or None)
    """
    # Check authentication first
    auth_error = _check_user_authentication(request)
    if auth_error:
        return False, None, auth_error

    # Get the project
    try:
        project = Project.objects.get(sodar_uuid=project_uuid)
    except Project.DoesNotExist:
        return False, None, JsonResponse({"error": "Project not found"}, status=404)

    # Check if user has any role in the project (guest or higher)
    has_role = RoleAssignment.objects.filter(project=project, user=request.user).exists()

    # Superusers always have access
    if request.user.is_superuser or has_role:
        return True, project, None
    else:
        return False, project, JsonResponse({"error": "Permission denied"}, status=403)


def _check_case_permission(request, case_info):
    """
    Check if the user has permission to access a case based on case_info.

    Args:
        request: Django request object
        case_info: Dictionary containing case information, must have 'name' key

    Returns:
        JsonResponse or None: Error response if permission denied, None if allowed
    """
    if not case_info or not case_info.get("name"):
        return None  # No case info provided, no permission check needed

    # Check authentication first
    auth_error = _check_user_authentication(request)
    if auth_error:
        return auth_error

    # Find the case by name and get its project
    try:
        case = Case.objects.select_related("project").get(name=case_info["name"])
    except Case.DoesNotExist:
        return JsonResponse({"error": "Case not found"}, status=404)
    except Case.MultipleObjectsReturned:
        return JsonResponse({"error": "Multiple cases found with the same name"}, status=400)

    # Check if user has any role in the case's project
    has_role = RoleAssignment.objects.filter(project=case.project, user=request.user).exists()

    # Superusers always have access
    if request.user.is_superuser or has_role:
        return None  # Permission granted
    else:
        return JsonResponse(
            {"error": "Permission denied - no access to case's project"}, status=403
        )


def _create_pdf_styles():
    """Create and return standardized PDF styles for use across all PDF generation functions."""
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
    )

    heading2_style = ParagraphStyle(
        "CustomHeading2",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=12,
        fontName="Helvetica-Bold",
    )

    heading3_style = ParagraphStyle(
        "CustomHeading3",
        parent=styles["Heading3"],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=8,
        fontName="Helvetica-Bold",
    )

    return styles, title_style, heading2_style, heading3_style


def _create_pdf_table(table_data, col_widths, font_size=9, padding=None):
    """
    Create a standardized PDF table with consistent styling.

    Args:
        table_data: List of lists containing table data (first row is header)
        col_widths: List of column widths in inches
        font_size: Font size for the table (default: 9)
        padding: Dict with padding values, defaults to standard padding

    Returns:
        Table object with applied styling
    """
    if padding is None:
        padding = {"left": 8, "right": 8, "top": 6, "bottom": 6}

    table = Table(table_data, colWidths=[w * inch for w in col_widths])
    table.setStyle(
        TableStyle(
            [
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),  # Header row
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),  # Data rows
                ("FONTSIZE", (0, 0), (-1, -1), font_size),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), padding["left"]),
                ("RIGHTPADDING", (0, 0), (-1, -1), padding["right"]),
                ("TOPPADDING", (0, 0), (-1, -1), padding["top"]),
                ("BOTTOMPADDING", (0, 0), (-1, -1), padding["bottom"]),
            ]
        )
    )
    return table


def _create_pdf_document_base(title="Applied Filter Settings"):
    """
    Create a PDF document with standard setup.

    Args:
        title: The title to set for the PDF document

    Returns:
        tuple: (buffer, doc, story, styles, title_style, heading2_style, heading3_style)
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=72,
        bottomMargin=18,
        title=title,
    )
    story = []
    styles, title_style, heading2_style, heading3_style = _create_pdf_styles()

    return buffer, doc, story, styles, title_style, heading2_style, heading3_style


def _create_metadata_table(meta_data):
    """
    Create a standardized metadata table.

    Args:
        meta_data: List of [key, value] pairs

    Returns:
        Table object with metadata styling
    """
    if not meta_data:
        return None

    table = _create_pdf_table(
        meta_data,
        [3, 4],  # Column widths
        font_size=10,
        padding={"left": 8, "right": 8, "top": 6, "bottom": 6},
    )
    return table


def generate_preset_pdf_directly(target_presetset, presetset_display_label, base_filename):
    """
    Generate PDF directly for preset settings using reportlab.

    Args:
        target_presetset: The PresetSet object
        presetset_display_label (str): Display label for the preset set
        base_filename (str): Base filename without extension

    Returns:
        tuple: (success: bool, pdf_content: bytes or None, error_msg: str or None)
    """
    try:
        # Create PDF document with proper title
        pdf_title = base_filename if base_filename else f"Query Presets - {presetset_display_label}"

        # Create PDF content using helper function
        buffer, doc, story, styles, title_style, heading2_style, heading3_style = (
            _create_pdf_document_base(pdf_title)
        )

        # Add title
        title = Paragraph(f"Query Presets - {presetset_display_label}", title_style)
        story.append(title)
        story.append(Spacer(1, 12))

        # Get quickpresets
        quickpresets = QuickPresets.objects.filter(presetset=target_presetset)

        if not quickpresets.exists():
            story.append(Paragraph("No quick presets found for this preset set.", styles["Normal"]))
        else:
            # Add each quickpreset as a section
            for quickpreset in quickpresets:
                # Add preset label as heading
                preset_label = quickpreset.label or "Unnamed Preset"
                story.append(Paragraph(preset_label, heading2_style))
                story.append(Spacer(1, 12))

                # Add inheritance if available
                if quickpreset.inheritance:
                    story.append(Paragraph("Inheritance", heading3_style))
                    story.append(Paragraph(str(quickpreset.inheritance), styles["Normal"]))
                    story.append(Spacer(1, 12))

                # Add preset categories
                preset_models = [
                    (ChromosomePresets, "chromosome", "Chromosomes"),
                    (QualityPresets, "quality", "Quality"),
                    (ImpactPresets, "impact", "Impact"),
                    (FrequencyPresets, "frequency", "Frequency"),
                    (FlagsEtcPresets, "flagsetc", "Flags"),
                ]

                for model_class, category_key, category_title in preset_models:
                    preset_obj = getattr(quickpreset, category_key, None)
                    if preset_obj:
                        preset_obj_label = preset_obj.label or f"Unnamed {category_title}"
                        story.append(
                            Paragraph(f'{category_title}: "{preset_obj_label}"', heading3_style)
                        )

                        # Get preset data (simplified for PDF)
                        preset_dict = model_to_dict(
                            preset_obj, exclude=["id", "date_created", "date_modified"]
                        )
                        cleaned_data = _clean_preset_dict(preset_dict)

                        if cleaned_data:
                            # Create a simple table with key-value pairs
                            table_data = []
                            for key, value in cleaned_data.items():
                                if key != "label":  # Skip label as it's used as heading
                                    if isinstance(value, list):
                                        if not value:  # Empty list
                                            value_str = "(empty)"
                                        elif (
                                            key == "effects" or len(value) > 5
                                        ):  # Show all effects as bullets, or if list is long
                                            value_str = "\n".join([f"• {str(v)}" for v in value])
                                        elif key in [
                                            "genomic_region",
                                            "gene_allowlist",
                                        ]:  # Always comma-separated for these fields
                                            value_str = ", ".join(str(v) for v in value)
                                        else:
                                            value_str = ", ".join(str(v) for v in value)
                                    elif isinstance(value, dict):
                                        value_str = f"({len(value)} settings)"
                                    else:
                                        value_str = str(value)

                                    table_data.append([str(key), value_str])

                            if table_data:
                                preset_table = Table(table_data, colWidths=[3 * inch, 4 * inch])
                                preset_table.setStyle(
                                    TableStyle(
                                        [
                                            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                                            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                                            ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                                            ("FONTSIZE", (0, 0), (-1, -1), 9),
                                            ("GRID", (0, 0), (-1, -1), 1, colors.black),
                                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                                            ("TOPPADDING", (0, 0), (-1, -1), 4),
                                            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                                        ]
                                    )
                                )
                                story.append(preset_table)
                        else:
                            story.append(Paragraph("No additional configuration", styles["Normal"]))

                        story.append(Spacer(1, 12))

                story.append(Spacer(1, 20))  # Extra space between presets

        # Build the PDF
        doc.build(story)
        buffer.seek(0)
        pdf_content = buffer.getvalue()
        buffer.close()

        return True, pdf_content, None

    except Exception as e:
        logger.error(f"Error generating preset PDF with reportlab: {str(e)}")
        return False, None, "PDF generation error occurred."


def generate_pdf_directly(filter_settings, case_info, request, base_filename):
    """
    Generate PDF directly using reportlab without DOCX intermediate step.

    Args:
        filter_settings (dict): The filter settings data
        case_info (dict): Information about the case
        request: The HTTP request object
        base_filename (str): Base filename without extension

    Returns:
        tuple: (success: bool, pdf_content: bytes or None, error_msg: str or None)
    """
    try:
        # Create PDF document with proper title
        pdf_title = base_filename if base_filename else "Applied Filter Settings"

        # Create PDF content using helper function
        buffer, doc, story, styles, title_style, heading2_style, heading3_style = (
            _create_pdf_document_base(pdf_title)
        )

        # Add title
        title = Paragraph("Applied Filter Settings", title_style)
        story.append(title)
        story.append(Spacer(1, 12))

        # Add metadata
        meta_data = []
        export_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        meta_data.append(["Export Date:", export_date])

        if case_info and case_info.get("name"):
            meta_data.append(["Case:", case_info["name"]])

        if request.user and hasattr(request.user, "username"):
            meta_data.append(["User:", request.user.username])

        if "database" in filter_settings:
            db_value = filter_settings["database"]
            if db_value == "refseq":
                display_value = "RefSeq"
            elif db_value == "ensembl":
                display_value = "EnsEMBL"
            else:
                display_value = str(db_value)
            meta_data.append(["Transcript Database:", display_value])

        # Create metadata table
        if meta_data:
            meta_table = Table(meta_data, colWidths=[3 * inch, 4 * inch])
            meta_table.setStyle(
                TableStyle(
                    [
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 10),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(meta_table)
            story.append(Spacer(1, 20))

        # Add complete filter settings
        if filter_settings:
            # Categorize filter settings using the same logic as DOCX export
            settings_dict = _categorize_filter_settings(filter_settings)

            # Add sections in the same order as DOCX export
            _add_pdf_categorized_sections(
                story,
                settings_dict,
                case_info,
                filter_settings,
                styles,
                heading2_style,
                heading3_style,
            )
        else:
            story.append(
                Paragraph("No filter settings data available to export.", styles["Normal"])
            )

        # Build the PDF
        doc.build(story)
        buffer.seek(0)
        pdf_content = buffer.getvalue()
        buffer.close()

        return True, pdf_content, None

    except Exception as e:
        logger.error(f"Error generating PDF with reportlab: {str(e)}")
        return False, None, "PDF generation error occurred."


def _add_pdf_categorized_sections(
    story, settings_dict, case_info, filter_settings, styles, heading2_style, heading3_style
):
    """Add categorized settings sections to the PDF story."""
    # Follow the same order as DOCX export
    # 1. Genotype tab
    if "genotype" in settings_dict["non_frequency"]:
        _add_pdf_genotype_section(
            story, settings_dict["non_frequency"]["genotype"], case_info, styles, heading2_style
        )

    # 2. Frequency tab
    if settings_dict["frequency"]:
        _add_pdf_frequency_section(story, settings_dict["frequency"], styles, heading2_style)

    # 3. Prioritization tab
    if settings_dict["prioritization"]:
        _add_pdf_prioritization_section(
            story, settings_dict["prioritization"], styles, heading2_style
        )

    # 4. Variants & Effects tab (Consequence)
    if settings_dict["consequence"]:
        _add_pdf_consequence_section(story, settings_dict["consequence"], styles, heading2_style)

    # 5. Quality tab
    if "quality" in settings_dict["non_frequency"]:
        _add_pdf_quality_section(
            story, settings_dict["non_frequency"]["quality"], styles, heading2_style
        )

    # 6. ClinVar tab
    if settings_dict["clinvar"]:
        _add_pdf_clinvar_section(story, settings_dict["clinvar"], styles, heading2_style)

    # 7. Gene Lists & Regions tab - always show, even if empty
    _add_pdf_genes_regions_section(story, settings_dict["genes_regions"], styles, heading2_style)

    # 8. Flags & Comments tab
    if settings_dict["flag"]:
        _add_pdf_flags_section(story, settings_dict["flag"], styles, heading2_style, heading3_style)

    # Add any remaining sections
    remaining_sections = ["locus", "inheritance"]
    for section_key in remaining_sections:
        if section_key in settings_dict["non_frequency"]:
            section_title = section_key.replace("_", " ").title() + " Settings"
            _add_pdf_generic_section(
                story,
                settings_dict["non_frequency"][section_key],
                section_title,
                styles,
                heading2_style,
            )


def _add_pdf_generic_section(story, data, title, styles, heading2_style):
    """Add a generic section to PDF."""

    if not data:
        return

    story.append(Paragraph(title, heading2_style))
    story.append(Spacer(1, 12))

    if isinstance(data, dict):
        table_data = []
        for key, value in data.items():
            if value is not None and value != "" and value != []:
                formatted_key = key.replace("_", " ").title()
                if isinstance(value, list):
                    if not value:
                        value_str = "(empty)"
                    elif len(value) > 10:
                        value_str = (
                            f"{', '.join(str(v) for v in value[:10])}... ({len(value)} total)"
                        )
                    else:
                        value_str = ", ".join(str(v) for v in value)
                elif isinstance(value, bool):
                    value_str = "Yes" if value else "No"
                elif isinstance(value, dict):
                    value_str = f"({len(value)} settings)"
                else:
                    value_str = str(value)
                table_data.append([formatted_key, value_str])

        if table_data:
            table = Table(table_data, colWidths=[3 * inch, 4 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(table)
    else:
        story.append(Paragraph(str(data), styles["Normal"]))

    story.append(Spacer(1, 20))


def _add_pdf_genotype_section(story, data, case_info, styles, heading2_style):
    """Add genotype settings section to PDF with pedigree information."""

    if not isinstance(data, dict):
        return

    story.append(Paragraph("Genotype Settings", heading2_style))
    story.append(Spacer(1, 12))

    # Filter out non-individual keys
    individuals = [key for key in data.keys() if key not in ["filter_active", "inheritance"]]

    if individuals:
        # Build pedigree lookup if case_info is available
        pedigree_lookup = {}
        if case_info and case_info.get("pedigree"):
            for member in case_info["pedigree"]:
                clean_name = clean_individual_name(member["name"])
                pedigree_lookup[clean_name] = member

        # Create table with columns: Individual, Father, Mother, Sex, Affected, Genotype
        table_data = [["Individual", "Father", "Mother", "Sex", "Affected", "Genotype"]]

        for individual in individuals:
            clean_name = clean_individual_name(individual)

            # Get pedigree info if available
            pedigree_info = pedigree_lookup.get(clean_name, {})

            # Father
            father = pedigree_info.get("father", "")
            father_text = clean_individual_name(father) if father else "-"

            # Mother
            mother = pedigree_info.get("mother", "")
            mother_text = clean_individual_name(mother) if mother else "-"

            # Sex (convert numeric to text)
            sex = pedigree_info.get("sex", "")
            if sex == 1:
                sex_text = "male"
            elif sex == 2:
                sex_text = "female"
            else:
                sex_text = "-"

            # Affected status (convert numeric to text)
            affected = pedigree_info.get("affected", "")
            if affected == 1:
                affected_text = "unaffected"
            elif affected == 2:
                affected_text = "affected"
            else:
                affected_text = "-"

            # Genotype value
            if individual in data:
                if isinstance(data[individual], dict):
                    genotype_val = data[individual].get("gt", str(data[individual]))
                else:
                    genotype_val = str(data[individual])
            else:
                genotype_val = "-"

            table_data.append(
                [clean_name, father_text, mother_text, sex_text, affected_text, genotype_val]
            )

        if len(table_data) > 1:  # More than just header
            # Create genotype table with smaller font and padding for 6 columns
            table = _create_pdf_table(
                table_data,
                [1.4, 1, 1, 0.8, 1.2, 1.6],  # Column widths
                font_size=8,  # Smaller font for more columns
                padding={"left": 4, "right": 4, "top": 4, "bottom": 4},  # Smaller padding
            )
            story.append(table)

    # Add inheritance setting if present
    if "inheritance" in data:
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Inheritance: {data['inheritance']}", styles["Normal"]))

    story.append(Spacer(1, 20))


def _add_pdf_frequency_section(story, data, styles, heading2_style):
    """Add frequency settings section to PDF."""

    if not isinstance(data, dict):
        return

    story.append(Paragraph("Frequency Settings", heading2_style))
    story.append(Spacer(1, 12))

    # Filter databases that have at least one setting present
    present_databases = []
    for db in FREQUENCY_DATABASES:
        has_data = False
        for field in [
            db["enabled_field"],
            db["homozygous_field"],
            db["heterozygous_field"],
            db["hemizygous_field"],
            db["frequency_field"],
        ]:
            if field and field in data and data[field] is not None and data[field] != "":
                has_data = True
                break
        if has_data:
            present_databases.append(db)

    if present_databases:
        # Create table headers
        table_data = [
            ["Database", "Enabled", "Homozygous", "Heterozygous", "Hemizygous", "Frequency"]
        ]

        for db in present_databases:
            row = [db["name"]]

            # Enabled column
            enabled_val = data.get(db["enabled_field"], False)
            row.append("Yes" if enabled_val else "No")

            # Homozygous, heterozygous, hemizygous, frequency columns
            for field in [
                db["homozygous_field"],
                db["heterozygous_field"],
                db["hemizygous_field"],
                db["frequency_field"],
            ]:
                if field and field in data:
                    val = data[field]
                    row.append(str(val) if val is not None and val != "" else "-")
                else:
                    row.append("N/A")

            table_data.append(row)

        if len(table_data) > 1:  # More than just header
            table = Table(
                table_data,
                colWidths=[1.4 * inch, 0.8 * inch, 1 * inch, 1 * inch, 1 * inch, 1.8 * inch],
            )
            table.setStyle(
                TableStyle(
                    [
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(table)
    else:
        story.append(Paragraph("No frequency settings configured.", styles["Normal"]))

    story.append(Spacer(1, 20))


def _add_pdf_consequence_section(story, data, styles, heading2_style):
    """Add variants & effects settings section to PDF."""

    if not isinstance(data, dict):
        return

    story.append(Paragraph("Variants & Effects Settings", heading2_style))
    story.append(Spacer(1, 12))

    # Create a unified table with all consequence settings including effects
    table_settings = []

    # Add effects first (if present)
    if "effects" in data:
        effects = data["effects"]
        if effects:
            if isinstance(effects, list) and len(effects) > 5:
                # For long lists, show as bulleted format
                effects_str = "\n".join([f"• {str(effect)}" for effect in effects])
            else:
                # For shorter lists, show as comma-separated
                effects_str = ", ".join(str(effect) for effect in effects) if effects else "(none)"
        else:
            effects_str = "(none)"
        table_settings.append(["Effects", effects_str])

    # Add all other settings
    for key, value in data.items():
        if key != "effects" and value is not None and value != "" and value != []:
            formatted_key = key.replace("_", " ").title()
            if isinstance(value, bool):
                value_str = "Yes" if value else "No"
            elif isinstance(value, list):
                value_str = ", ".join(str(v) for v in value) if value else "(empty)"
            else:
                value_str = str(value)
            table_settings.append([formatted_key, value_str])

    if table_settings:
        table = Table(table_settings, colWidths=[3 * inch, 4 * inch])
        table.setStyle(
            TableStyle(
                [
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(table)

    story.append(Spacer(1, 20))


def _add_pdf_quality_section(story, data, styles, heading2_style):
    """Add quality settings section to PDF."""

    if not isinstance(data, dict):
        return

    story.append(Paragraph("Quality Settings", heading2_style))
    story.append(Spacer(1, 12))

    # Extract individual names (keys that aren't quality metric names)
    individuals = [
        key for key in data.keys() if key not in QUALITY_METRICS and key != "filter_active"
    ]

    if individuals:
        # Create table with individuals as rows and quality metrics as columns
        table_data = [["Individual"] + [metric.upper() for metric in QUALITY_METRICS]]

        for individual in individuals:
            row = [clean_individual_name(individual)]
            for metric in QUALITY_METRICS:
                if (
                    individual in data
                    and isinstance(data[individual], dict)
                    and metric in data[individual]
                ):
                    row.append(str(data[individual][metric]))
                else:
                    row.append("-")
            table_data.append(row)

        if len(table_data) > 1:  # More than just header
            # Standard 7-inch total width: Individual column + 6 quality metrics
            col_widths = [1.6 * inch] + [0.9 * inch] * len(QUALITY_METRICS)
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(
                TableStyle(
                    [
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(table)
    else:
        story.append(Paragraph("No quality settings configured.", styles["Normal"]))

    story.append(Spacer(1, 20))


def _add_pdf_prioritization_section(story, data, styles, heading2_style):
    """Add prioritization settings section to PDF."""

    story.append(Paragraph("Prioritization Settings", heading2_style))
    story.append(Spacer(1, 12))

    if not isinstance(data, dict):
        story.append(Paragraph("No prioritization settings configured.", styles["Normal"]))
        story.append(Spacer(1, 20))
        return

    # Create table with all prioritization settings
    table_data = []

    # Phenotype prioritization settings
    prio_enabled = data.get("prio_enabled", False) if "prio_enabled" in data else False
    if "prio_enabled" in data:
        table_data.append(
            ["Enable phenotype-based prioritization", "Yes" if prio_enabled else "No"]
        )

    # Algorithm setting - show "(empty)" if phenotype prioritization is disabled
    if "prio_algorithm" in data or "prio_enabled" in data:
        if prio_enabled and "prio_algorithm" in data:
            # Map algorithm codes to readable names
            algorithm_map = {
                "phenix": "Phenix",
                "phive": "Phive",
                "hiphive-human": "HiPhive (human only)",
                "hiphive-mouse": "HiPhive (human+mouse)",
                "hiphive": "HiPhive (human, mouse, fish, PPI)",
            }
            algorithm_label = algorithm_map.get(data["prio_algorithm"], data["prio_algorithm"])
        else:
            algorithm_label = "(empty)"
        table_data.append(["Phenotype similarity algorithm", algorithm_label])

    # HPO terms - show "(empty)" if phenotype prioritization is disabled
    if prio_enabled:
        # Show actual HPO terms if phenotype prioritization is enabled
        if "prio_hpo_terms" in data:
            hpo_terms = format_list_value(data["prio_hpo_terms"])
        else:
            hpo_terms = "(none)"
    else:
        # Show "(empty)" if phenotype prioritization is disabled
        hpo_terms = "(empty)"
    table_data.append(["HPO Terms", hpo_terms])

    # Pathogenicity prioritization settings
    patho_enabled = data.get("patho_enabled", False) if "patho_enabled" in data else False
    if "patho_enabled" in data:
        table_data.append(
            ["Enable pathogenicity-based prioritization", "Yes" if patho_enabled else "No"]
        )

    # Pathogenicity score - show "(empty)" if pathogenicity prioritization is disabled
    if "patho_score" in data or "patho_enabled" in data:
        if patho_enabled and "patho_score" in data:
            # Map score codes to readable names
            score_map = {"cadd": "CADD"}
            score_label = score_map.get(data["patho_score"], data["patho_score"])
        else:
            score_label = "(empty)"
        table_data.append(["Pathogenicity score", score_label])

    # Other prioritization settings
    if "gm_enabled" in data:
        table_data.append(
            ["Enable GestaltMatcher-based prioritization", "Yes" if data["gm_enabled"] else "No"]
        )

    if "pedia_enabled" in data:
        table_data.append(
            ["Enable PEDIA based prioritization", "Yes" if data["pedia_enabled"] else "No"]
        )

    # Add other settings that might be present
    handled_keys = {
        "prio_enabled",
        "prio_algorithm",
        "prio_hpo_terms",
        "patho_enabled",
        "patho_score",
        "gm_enabled",
        "pedia_enabled",
    }
    for key, value in data.items():
        if key not in handled_keys and value is not None and value != "" and value != []:
            formatted_key = key.replace("_", " ").title()
            if isinstance(value, bool):
                value_str = "Yes" if value else "No"
            elif isinstance(value, list):
                value_str = format_list_value(value)
            else:
                value_str = str(value)
            table_data.append([formatted_key, value_str])

    if table_data:
        table = Table(table_data, colWidths=[3 * inch, 4 * inch])
        table.setStyle(
            TableStyle(
                [
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(table)
    else:
        story.append(Paragraph("No prioritization settings configured.", styles["Normal"]))

    story.append(Spacer(1, 20))


def _add_pdf_clinvar_section(story, data, styles, heading2_style):
    """Add ClinVar settings section to PDF."""
    _add_pdf_generic_section(story, data, "ClinVar Settings", styles, heading2_style)


def _add_pdf_genes_regions_section(story, data, styles, heading2_style):
    """Add gene lists & regions settings section to PDF."""

    story.append(Paragraph("Gene Lists & Regions Settings", heading2_style))
    story.append(Spacer(1, 12))

    # Always create a table with both gene allowlist and genomic regions, even if empty
    table_data = [
        ["Setting", "Value"],  # Header
        ["Gene Allow List", ""],  # Will be filled below
        ["Genomic Regions", ""],  # Will be filled below
    ]

    # Get gene allowlist and format it
    gene_allowlist = data.get("gene_allowlist", []) if isinstance(data, dict) else []
    table_data[1][1] = format_list_value(gene_allowlist, max_items=10)

    # Get genomic regions and format them
    genomic_regions = data.get("genomic_region", []) if isinstance(data, dict) else []
    table_data[2][1] = format_list_value(genomic_regions, max_items=5)

    table = Table(table_data, colWidths=[3 * inch, 4 * inch])
    table.setStyle(
        TableStyle(
            [
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),  # Header row
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),  # Data rows
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 20))


def _add_pdf_flags_section(story, data, styles, heading2_style, heading3_style):
    """Add flags settings section to PDF matching FlagsPane.vue structure."""

    if not isinstance(data, dict):
        return

    story.append(Paragraph("Flags & Comments Settings", heading2_style))
    story.append(Spacer(1, 12))

    # Simple Flags section (first section in FlagsPane.vue)
    simple_flags = [
        {"id": "flag_bookmarked", "label": "bookmarked"},
        {"id": "flag_incidental", "label": "incidental finding"},
        {"id": "flag_candidate", "label": "candidate"},
        {"id": "flag_final_causative", "label": "reported"},
        {"id": "flag_for_validation", "label": "for validation"},
        {"id": "flag_no_disease_association", "label": "no disease association"},
        {"id": "flag_segregates", "label": "segregates"},
        {"id": "flag_doesnt_segregate", "label": "does not segregate"},
        {"id": "flag_simple_empty", "label": "no simple flag"},
    ]

    # Create Simple Flags matrix table
    has_simple_flags = any(flag["id"] in data for flag in simple_flags)

    if has_simple_flags:
        # Create a heading for simple flags
        story.append(Paragraph("Simple Flags", heading3_style))

        # Create simple flags matrix: each flag gets a row with Yes/No
        simple_table_data = [["Flag", "Enabled"]]  # Header

        for flag in simple_flags:
            if flag["id"] in data:  # Only show flags that exist in data
                enabled = "Yes" if data[flag["id"]] else "No"
                simple_table_data.append([flag["label"], enabled])

        if len(simple_table_data) > 1:  # More than just header
            simple_table = Table(simple_table_data, colWidths=[4 * inch, 3 * inch])
            simple_table.setStyle(
                TableStyle(
                    [
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),  # Header row
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),  # Data rows
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(simple_table)
            story.append(Spacer(1, 16))

    # Create Flag Categories matrix table
    flag_categories = [
        {"id": "flag_visual", "label": "Visual"},
        {"id": "flag_validation", "label": "Validation"},
        {"id": "flag_phenotype_match", "label": "Phenotype Match"},
        {"id": "flag_molecular", "label": "Molecular Impact"},
        {"id": "flag_summary", "label": "Summary"},
    ]

    flag_values = [
        {"id": "positive", "label": "Positive"},
        {"id": "uncertain", "label": "Uncertain"},
        {"id": "negative", "label": "Negative"},
        {"id": "empty", "label": "Empty"},
    ]

    # Check if we have any category flag data
    has_category_flags = any(
        f"{category['id']}_{value['id']}" in data
        for category in flag_categories
        for value in flag_values
    )

    if has_category_flags:
        # Create a heading for flag categories
        story.append(Paragraph("Flag Categories", heading3_style))

        # Create category flags matrix: categories as rows, values as columns
        header_row = ["Flag Category"] + [value["label"] for value in flag_values]
        category_table_data = [header_row]

        for category in flag_categories:
            # Check if this category has any data
            has_category_data = any(
                f"{category['id']}_{value['id']}" in data for value in flag_values
            )

            if has_category_data:
                row = [category["label"]]
                for value in flag_values:
                    field_name = f"{category['id']}_{value['id']}"
                    if field_name in data:
                        enabled = "Yes" if data[field_name] else "No"
                    else:
                        enabled = "No"
                    row.append(enabled)
                category_table_data.append(row)

        if len(category_table_data) > 1:  # More than just header
            # Column widths: Category name gets more space, values get equal smaller columns
            col_widths = [2.2 * inch] + [1.2 * inch] * len(flag_values)
            category_table = Table(category_table_data, colWidths=col_widths)
            category_table.setStyle(
                TableStyle(
                    [
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),  # Header row
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),  # Data rows
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(category_table)

    # If no flag data at all
    if not has_simple_flags and not has_category_flags:
        story.append(Paragraph("No flag settings configured.", styles["Normal"]))

    story.append(Spacer(1, 20))


# Constants for field names and mappings
QUALITY_METRICS = ["ab", "ad", "dp_het", "dp_hom", "gq", "fail"]

FREQUENCY_DATABASES = [
    {
        "name": "1000 Genomes",
        "prefix": "thousand_genomes",
        "enabled_field": "thousand_genomes_enabled",
        "homozygous_field": "thousand_genomes_homozygous",
        "heterozygous_field": "thousand_genomes_heterozygous",
        "hemizygous_field": "thousand_genomes_hemizygous",
        "frequency_field": "thousand_genomes_frequency",
    },
    {
        "name": "ExAC",
        "prefix": "exac",
        "enabled_field": "exac_enabled",
        "homozygous_field": "exac_homozygous",
        "heterozygous_field": "exac_heterozygous",
        "hemizygous_field": "exac_hemizygous",
        "frequency_field": "exac_frequency",
    },
    {
        "name": "gnomAD exomes",
        "prefix": "gnomad_exomes",
        "enabled_field": "gnomad_exomes_enabled",
        "homozygous_field": "gnomad_exomes_homozygous",
        "heterozygous_field": "gnomad_exomes_heterozygous",
        "hemizygous_field": "gnomad_exomes_hemizygous",
        "frequency_field": "gnomad_exomes_frequency",
    },
    {
        "name": "gnomAD genomes",
        "prefix": "gnomad_genomes",
        "enabled_field": "gnomad_genomes_enabled",
        "homozygous_field": "gnomad_genomes_homozygous",
        "heterozygous_field": "gnomad_genomes_heterozygous",
        "hemizygous_field": "gnomad_genomes_hemizygous",
        "frequency_field": "gnomad_genomes_frequency",
    },
    {
        "name": "in-house DB",
        "prefix": "inhouse",
        "enabled_field": "inhouse_enabled",
        "homozygous_field": "inhouse_homozygous",
        "heterozygous_field": "inhouse_heterozygous",
        "hemizygous_field": "inhouse_hemizygous",
        "frequency_field": "inhouse_carriers",  # Special case: carriers instead of frequency
    },
    {
        "name": "mtDB",
        "prefix": "mtdb",
        "enabled_field": "mtdb_enabled",
        "homozygous_field": "mtdb_count",  # Special case: uses count instead of homozygous
        "heterozygous_field": None,  # N/A for mtDB
        "hemizygous_field": None,  # N/A for mtDB
        "frequency_field": "mtdb_frequency",
    },
    {
        "name": "HelixMTdb",
        "prefix": "helixmtdb",
        "enabled_field": "helixmtdb_enabled",
        "homozygous_field": "helixmtdb_hom_count",
        "heterozygous_field": "helixmtdb_het_count",
        "hemizygous_field": None,  # N/A for HelixMTdb
        "frequency_field": "helixmtdb_frequency",
    },
    {
        "name": "MITOMAP",
        "prefix": "mitomap",
        "enabled_field": "mitomap_enabled",
        "homozygous_field": "mitomap_count",  # Special case: uses count instead of homozygous
        "heterozygous_field": None,  # N/A for MITOMAP
        "hemizygous_field": None,  # N/A for MITOMAP
        "frequency_field": "mitomap_frequency",
    },
]

SIMPLE_FLAGS = [
    {"id": "flag_bookmarked", "label": "bookmarked"},
    {"id": "flag_incidental", "label": "incidental finding"},
    {"id": "flag_candidate", "label": "candidate"},
    {"id": "flag_final_causative", "label": "reported"},
    {"id": "flag_for_validation", "label": "for validation"},
    {"id": "flag_no_disease_association", "label": "no disease association"},
    {"id": "flag_segregates", "label": "segregates"},
    {"id": "flag_doesnt_segregate", "label": "does not segregate"},
    {"id": "flag_simple_empty", "label": "no simple flag"},
]

FLAG_NAMES = [
    {"id": "flag_visual", "label": "Visual"},
    {"id": "flag_validation", "label": "Validation"},
    {"id": "flag_phenotype_match", "label": "Phenotype Match"},
    {"id": "flag_molecular", "label": "Molecular Impact"},
    {"id": "flag_summary", "label": "Summary"},
]

FLAG_VALUES = [
    {"id": "positive", "label": "positive"},
    {"id": "uncertain", "label": "uncertain"},
    {"id": "negative", "label": "negative"},
    {"id": "empty", "label": "empty"},
]

CLINVAR_INTERPRETATIONS = [
    {"id": "pathogenic", "label": "P5 - pathogenic"},
    {"id": "likely_pathogenic", "label": "LP4 - likely pathogenic"},
    {"id": "uncertain_significance", "label": "VUS3 - uncertain significance"},
    {"id": "likely_benign", "label": "LB2 - likely benign"},
    {"id": "benign", "label": "B1 - benign"},
]

PHENOTYPE_ALGORITHM_OPTIONS = {
    "phenix": "Phenix",
    "phive": "Phive",
    "hiphive-human": "HiPhive (human only)",
    "hiphive-mouse": "HiPhive (human+mouse)",
    "hiphive": "HiPhive (human, mouse, fish, PPI)",
}

PATHOGENICITY_SCORE_OPTIONS = {"cadd": "CADD"}

VARIANT_TYPE_MAPPINGS = {
    "var_type_snv": "SNV",
    "var_type_indel": "indel",
    "var_type_mnv": "MNV",
}

TRANSCRIPT_TYPE_MAPPINGS = {
    "transcripts_coding": "coding",
    "transcripts_noncoding": "non-coding",
}

# Field groupings for categorization
CONSEQUENCE_FIELDS = [
    "effects",
    "var_type_snv",
    "var_type_indel",
    "var_type_mnv",
    "transcripts_coding",
    "transcripts_noncoding",
    "max_exon_dist",
]

FLAG_FIELDS = [
    "flag_bookmarked",
    "flag_incidental",
    "flag_candidate",
    "flag_final_causative",
    "flag_for_validation",
    "flag_no_disease_association",
    "flag_segregates",
    "flag_doesnt_segregate",
    "flag_simple_empty",
    "flag_visual_positive",
    "flag_visual_uncertain",
    "flag_visual_negative",
    "flag_visual_empty",
    "flag_validation_positive",
    "flag_validation_uncertain",
    "flag_validation_negative",
    "flag_validation_empty",
    "flag_phenotype_match_positive",
    "flag_phenotype_match_uncertain",
    "flag_phenotype_match_negative",
    "flag_phenotype_match_empty",
    "flag_molecular_positive",
    "flag_molecular_uncertain",
    "flag_molecular_negative",
    "flag_molecular_empty",
    "flag_summary_positive",
    "flag_summary_uncertain",
    "flag_summary_negative",
    "flag_summary_empty",
]

CLINVAR_FIELDS = [
    "require_in_clinvar",
    "clinvar_include_pathogenic",
    "clinvar_include_likely_pathogenic",
    "clinvar_include_uncertain_significance",
    "clinvar_include_likely_benign",
    "clinvar_include_benign",
    "clinvar_exclude_conflicting",
]

PRIORITIZATION_FIELDS = [
    "prio_enabled",
    "prio_algorithm",
    "prio_hpo_terms",
    "patho_enabled",
    "patho_score",
    "gm_enabled",
    "pedia_enabled",
]

GENES_REGIONS_FIELDS = ["gene_allowlist", "genomic_region"]

FREQUENCY_PREFIXES = [
    "thousand_genomes_",
    "exac_",
    "gnomad_exomes_",
    "gnomad_genomes_",
    "inhouse_",
    "mtdb_",
    "helixmtdb_",
    "mitomap_",
]


def clean_individual_name(name):
    """Clean individual names by removing sequencing run suffixes."""
    # Remove patterns like "-N1-DNA1-WGS1", "-N2-DNA2-WES1", etc.
    cleaned = re.sub(r"-N\d+-DNA\d+-(WGS|WES)\d+$", "", name)
    return cleaned


def format_display_name(field_name, custom_mappings=None):
    """
    Convert field names to human-readable display names.

    Args:
        field_name (str): The original field name to convert.
        custom_mappings (dict, optional): Custom mapping overrides for specific field names.

    Returns:
        str: Human-readable display name.
    """
    if custom_mappings and field_name in custom_mappings:
        return custom_mappings[field_name]

    # Standard field name to display name conversion
    display_name = field_name.replace("_", " ").title()
    return display_name


def format_boolean_value(value):
    """
    Format boolean values consistently.

    Args:
        value: The value to format (should be boolean).

    Returns:
        str: "Yes" for True, "No" for False.
    """
    return "Yes" if value else "No"


def format_list_value(value, max_items=5):
    """
    Format list values with truncation if too long.

    Args:
        value: The list value to format.
        max_items (int): Maximum number of items to display before truncating.

    Returns:
        str: Formatted string representation of the list.
    """
    if not value or len(value) == 0:
        return "(none)"

    if not isinstance(value, list):
        return str(value)

    if len(value) <= max_items:
        return ", ".join(str(item) for item in value)
    else:
        displayed_items = value[:max_items]
        items_str = ", ".join(str(item) for item in displayed_items)
        return f"{items_str} ... ({len(value)} total items)"


def create_simple_table(doc, title, rows_data, headers=None):
    """
    Create a simple two-column table with optional custom headers.

    Args:
        doc: The docx Document object.
        title (str): The section title for the table.
        rows_data (list): List of tuples containing (setting, value) pairs.
        headers (list, optional): Custom headers for the table columns.
    """
    if not rows_data:
        return

    if headers is None:
        headers = ["Setting", "Value"]

    doc.add_heading(title, level=2)

    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.style = "Table Grid"

    # Header
    header_row = table.rows[0]
    header_row.cells[0].text = headers[0]
    header_row.cells[1].text = headers[1]
    make_table_headers_bold(table)

    # Add rows
    for row_idx, (setting, value) in enumerate(rows_data):
        row = table.rows[row_idx + 1]
        row.cells[0].text = setting
        row.cells[1].text = value


def make_table_headers_bold(table, header_row_index=0):
    """
    Make all cells in the header row bold.

    Args:
        table: The docx table object.
        header_row_index (int): Index of the header row (default: 0).
    """
    header_row = table.rows[header_row_index]
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def add_key_value_table(doc, title, data):
    """
    Add a section with key-value pairs to the document using specialized formatters.

    Args:
        doc: The docx Document object.
        title (str): The section title.
        data: The data to be formatted and added to the document.
    """
    if not data or (isinstance(data, dict) and not data):
        return

    try:
        # Add section title
        doc.add_heading(title, level=1)

        # Special formatting for different section types
        if title == "Quality Settings":
            add_quality_settings_table(doc, data)
        elif title == "Genotype Settings":
            add_genotype_settings_table(doc, data)
        elif title == "Variants & Effects Settings":
            add_consequence_settings_section(doc, data)
        elif title == "Frequency Settings":
            add_frequency_settings_table(doc, data)
        elif "Flag" in title:
            add_flags_settings_table(doc, data)
        elif "ClinVar" in title:
            add_clinvar_settings_table(doc, data)
        elif "Gene Lists & Regions Settings" in title:
            add_genes_regions_settings_table(doc, data)
        elif "Other Settings" in title:
            add_other_settings_table(doc, data)
        elif "Prioritization" in title:
            add_prioritization_settings_table(doc, data)
        else:
            # Default formatting for other sections
            add_generic_section(doc, data)

        # Add some spacing
        doc.add_paragraph()
    except Exception as e:
        logger.exception(f"Error adding section {title}: {str(e)}")
        # Add at least the title if formatting fails
        try:
            doc.add_heading(f"{title} (formatting error)", level=1)
            doc.add_paragraph(f"Raw data: {str(data)}")
        except Exception as inner_e:
            logger.exception(f"Error adding fallback section for {title}: {str(inner_e)}")


def add_quality_settings_table(doc, data):
    """Add quality settings as a table with individuals as rows and quality metrics as columns."""
    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    # Extract individual names (keys that aren't quality metric names)
    individuals = [
        key for key in data.keys() if key not in QUALITY_METRICS and key != "filter_active"
    ]

    if not individuals:
        add_generic_section(doc, data)
        return

    # Create table: rows = individuals + 1 header, cols = quality metrics + 1 individual name
    table = doc.add_table(rows=len(individuals) + 1, cols=len(QUALITY_METRICS) + 1)
    table.style = "Table Grid"

    # Add header row
    header_row = table.rows[0]
    header_row.cells[0].text = "Individual"
    for i, metric in enumerate(QUALITY_METRICS):
        header_row.cells[i + 1].text = metric.upper()
    make_table_headers_bold(table)

    # Add individual rows
    for row_idx, individual in enumerate(individuals):
        row = table.rows[row_idx + 1]
        row.cells[0].text = clean_individual_name(individual)  # Clean name, no bold

        for col_idx, metric in enumerate(QUALITY_METRICS):
            if (
                individual in data
                and isinstance(data[individual], dict)
                and metric in data[individual]
            ):
                row.cells[col_idx + 1].text = str(data[individual][metric])
            else:
                row.cells[col_idx + 1].text = "-"


def add_genotype_settings_table_with_pedigree(doc, title, data, case_info):
    """Add genotype settings as a detailed table with pedigree information."""
    # Add section title
    doc.add_heading(title, level=1)

    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    # Filter out non-individual keys
    individuals = [key for key in data.keys() if key not in ["filter_active", "inheritance"]]

    if not individuals:
        add_generic_section(doc, data)
        return

    # Build pedigree lookup if case_info is available
    pedigree_lookup = {}
    if case_info and case_info.get("pedigree"):
        for member in case_info["pedigree"]:
            clean_name = clean_individual_name(member["name"])
            pedigree_lookup[clean_name] = member

    # Create table with columns: Individual, Father, Mother, Sex, Affected, Genotype
    table = doc.add_table(rows=len(individuals) + 1, cols=6)
    table.style = "Table Grid"

    # Add header row
    header_row = table.rows[0]
    headers = ["Individual", "Father", "Mother", "Sex", "Affected", "Genotype"]
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
    make_table_headers_bold(table)

    # Add individual rows
    for row_idx, individual in enumerate(individuals):
        row = table.rows[row_idx + 1]
        clean_name = clean_individual_name(individual)
        row.cells[0].text = clean_name  # Clean individual name

        # Get pedigree info if available
        pedigree_info = pedigree_lookup.get(clean_name, {})

        # Father
        father = pedigree_info.get("father", "")
        row.cells[1].text = clean_individual_name(father) if father else "-"

        # Mother
        mother = pedigree_info.get("mother", "")
        row.cells[2].text = clean_individual_name(mother) if mother else "-"

        # Sex (convert numeric to text)
        sex = pedigree_info.get("sex", "")
        if sex == 1:
            sex_text = "male"
        elif sex == 2:
            sex_text = "female"
        else:
            sex_text = "-"
        row.cells[3].text = sex_text

        # Affected status (convert numeric to text)
        affected = pedigree_info.get("affected", "")
        if affected == 1:
            affected_text = "unaffected"
        elif affected == 2:
            affected_text = "affected"
        else:
            affected_text = "-"
        row.cells[4].text = affected_text

        # Genotype value
        if individual in data:
            if isinstance(data[individual], dict):
                # If it's a dict, show the genotype value or the whole dict
                genotype_val = data[individual].get("gt", str(data[individual]))
            else:
                genotype_val = str(data[individual])
            row.cells[5].text = genotype_val
        else:
            row.cells[5].text = "-"

    # Add inheritance setting if present
    if "inheritance" in data:
        doc.add_heading("Inheritance", level=2)
        p = doc.add_paragraph()
        p.add_run(str(data["inheritance"]))

    # Add some spacing
    doc.add_paragraph()


def add_genotype_settings_table(doc, data):
    """Add genotype settings as a detailed table resembling GenotypePane.vue structure."""
    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    # Filter out non-individual keys
    individuals = [key for key in data.keys() if key not in ["filter_active", "inheritance"]]

    if not individuals:
        add_generic_section(doc, data)
        return

    # Create table with columns: Individual, Father, Mother, Sex, Affected, Genotype
    # (removed the # column as requested)
    table = doc.add_table(rows=len(individuals) + 1, cols=6)
    table.style = "Table Grid"

    # Add header row
    header_row = table.rows[0]
    headers = ["Individual", "Father", "Mother", "Sex", "Affected", "Genotype"]
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
    make_table_headers_bold(table)

    # Add individual rows
    for row_idx, individual in enumerate(individuals):
        row = table.rows[row_idx + 1]
        row.cells[0].text = clean_individual_name(individual)  # Clean individual name
        row.cells[1].text = "-"  # Father (not available in current data)
        row.cells[2].text = "-"  # Mother (not available in current data)
        row.cells[3].text = "-"  # Sex (not available in current data)
        row.cells[4].text = "-"  # Affected (not available in current data)

        # Genotype value
        if individual in data:
            if isinstance(data[individual], dict):
                # If it's a dict, show the genotype value or the whole dict
                genotype_val = data[individual].get("gt", str(data[individual]))
            else:
                genotype_val = str(data[individual])
            row.cells[5].text = genotype_val
        else:
            row.cells[5].text = "-"

    # Add inheritance setting if present
    if "inheritance" in data:
        doc.add_heading("Inheritance", level=2)
        p = doc.add_paragraph()
        p.add_run(str(data["inheritance"]))


def add_consequence_settings_section(doc, data):
    """Add consequence settings with effects as list and other settings."""
    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    # Handle effects as a list
    if "effects" in data:
        doc.add_heading("Effects", level=2)
        # Create bullet list for effects
        for effect in data["effects"]:
            bullet_p = doc.add_paragraph()
            bullet_p.style = "List Bullet"
            bullet_p.add_run(str(effect))

    # Handle variant types subsection
    var_type_fields = ["var_type_snv", "var_type_indel", "var_type_mnv"]
    var_type_data = {field: data.get(field) for field in var_type_fields if field in data}

    if var_type_data:
        doc.add_heading("Variant Types", level=2)

        # Create a simple table for variant types
        table = doc.add_table(rows=len(var_type_data) + 1, cols=2)
        table.style = "Table Grid"

        # Header
        header_row = table.rows[0]
        header_row.cells[0].text = "Variant Type"
        header_row.cells[1].text = "Enabled"
        make_table_headers_bold(table)

        # Data rows
        for i, (field, value) in enumerate(var_type_data.items()):
            row = table.rows[i + 1]
            # Convert field name to display name using mappings
            display_name = VARIANT_TYPE_MAPPINGS.get(field, field.replace("var_type_", "").upper())
            row.cells[0].text = display_name
            row.cells[1].text = format_boolean_value(value)

    # Handle transcript types subsection
    transcript_type_fields = ["transcripts_coding", "transcripts_noncoding"]
    transcript_type_data = {
        field: data.get(field) for field in transcript_type_fields if field in data
    }

    if transcript_type_data:
        doc.add_heading("Transcript Types", level=2)

        # Create a simple table for transcript types
        table = doc.add_table(rows=len(transcript_type_data) + 1, cols=2)
        table.style = "Table Grid"

        # Header
        header_row = table.rows[0]
        header_row.cells[0].text = "Transcript Type"
        header_row.cells[1].text = "Enabled"
        make_table_headers_bold(table)

        # Data rows
        for i, (field, value) in enumerate(transcript_type_data.items()):
            row = table.rows[i + 1]
            # Convert field name to display name using mappings
            display_name = TRANSCRIPT_TYPE_MAPPINGS.get(
                field, field.replace("transcripts_", "").title()
            )
            row.cells[0].text = display_name
            row.cells[1].text = format_boolean_value(value)

    # Handle max_exon_dist if present and not null
    if "max_exon_dist" in data and data["max_exon_dist"] is not None:
        doc.add_heading("Maximum Exon Distance", level=2)
        p = doc.add_paragraph()
        p.add_run(str(data["max_exon_dist"]) + " bp")

    # Handle other consequence settings
    handled_keys = [
        "effects",
        "var_type_snv",
        "var_type_indel",
        "var_type_mnv",
        "transcripts_coding",
        "transcripts_noncoding",
        "max_exon_dist",
    ]
    other_keys = [key for key in data.keys() if key not in handled_keys]

    if other_keys:
        for key in other_keys:
            if data[key] is not None and data[key] != "" and data[key] != []:
                p = doc.add_paragraph()
                p.add_run(f"{format_display_name(key)}: ").bold = True
                p.add_run(str(data[key]))


def add_frequency_settings_table(doc, data):
    """Add frequency settings as a table matching the FrequencyPane.vue format."""
    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    # Filter databases that have at least one setting present
    present_databases = []
    for db in FREQUENCY_DATABASES:
        has_data = False
        for field in [
            db["enabled_field"],
            db["homozygous_field"],
            db["heterozygous_field"],
            db["hemizygous_field"],
            db["frequency_field"],
        ]:
            if field and field in data and data[field] is not None and data[field] != "":
                has_data = True
                break
        if has_data:
            present_databases.append(db)

    if not present_databases:
        doc.add_paragraph("No frequency settings configured.")
        return

    # Create table: databases as rows + header, 6 columns (database, enabled, hom, het, hemi, freq)
    table = doc.add_table(rows=len(present_databases) + 1, cols=6)
    table.style = "Table Grid"

    # Add header row
    header_row = table.rows[0]
    headers = [
        "Database",
        "Enabled",
        "Homozygous Count",
        "Heterozygous Count",
        "Hemizygous Count",
        "Frequency/Carriers",
    ]
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
    make_table_headers_bold(table)

    # Add database rows
    for row_idx, db in enumerate(present_databases):
        row = table.rows[row_idx + 1]

        # Database name column
        row.cells[0].text = db["name"]

        # Enabled column
        enabled_val = data.get(db["enabled_field"], False)
        row.cells[1].text = format_boolean_value(enabled_val)

        # Homozygous count column
        if db["homozygous_field"] and db["homozygous_field"] in data:
            hom_val = data[db["homozygous_field"]]
            row.cells[2].text = str(hom_val) if hom_val is not None and hom_val != "" else "-"
        else:
            row.cells[2].text = "N/A"

        # Heterozygous count column
        if db["heterozygous_field"] and db["heterozygous_field"] in data:
            het_val = data[db["heterozygous_field"]]
            row.cells[3].text = str(het_val) if het_val is not None and het_val != "" else "-"
        else:
            row.cells[3].text = "N/A"

        # Hemizygous count column
        if db["hemizygous_field"] and db["hemizygous_field"] in data:
            hemi_val = data[db["hemizygous_field"]]
            row.cells[4].text = str(hemi_val) if hemi_val is not None and hemi_val != "" else "-"
        else:
            row.cells[4].text = "N/A"

        # Frequency/Carriers column
        if db["frequency_field"] and db["frequency_field"] in data:
            freq_val = data[db["frequency_field"]]
            row.cells[5].text = str(freq_val) if freq_val is not None and freq_val != "" else "-"
        else:
            row.cells[5].text = "-"


def add_flags_settings_table(doc, data):
    """Add flags settings as a table resembling the FlagsPane.vue structure."""
    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    # Check if we have any flag data
    has_simple_flags = any(flag["id"] in data and data[flag["id"]] for flag in SIMPLE_FLAGS)
    has_category_flags = any(
        f"{name['id']}_{value['id']}" in data and data[f"{name['id']}_{value['id']}"]
        for name in FLAG_NAMES
        for value in FLAG_VALUES
    )

    if not has_simple_flags and not has_category_flags:
        doc.add_paragraph("No flag settings configured.")
        return

    # Add Simple Flags section
    if has_simple_flags:
        doc.add_heading("Simple Flags", level=2)

        # Create table for simple flags
        enabled_simple_flags = [
            flag for flag in SIMPLE_FLAGS if flag["id"] in data and data[flag["id"]]
        ]

        if enabled_simple_flags:
            table = doc.add_table(rows=len(enabled_simple_flags) + 1, cols=2)
            table.style = "Table Grid"

            # Header
            header_row = table.rows[0]
            header_row.cells[0].text = "Flag Type"
            header_row.cells[1].text = "Enabled"
            make_table_headers_bold(table)

            # Add enabled flags
            for row_idx, flag in enumerate(enabled_simple_flags):
                row = table.rows[row_idx + 1]
                row.cells[0].text = flag["label"]
                row.cells[1].text = "Yes"

    # Add Category Flags section
    if has_category_flags:
        doc.add_heading("Flag Categories", level=2)

        # Create table for category flags: Flag Category | positive | uncertain | negative | empty
        table = doc.add_table(rows=len(FLAG_NAMES) + 1, cols=5)
        table.style = "Table Grid"

        # Header row
        header_row = table.rows[0]
        header_row.cells[0].text = "Flag Category"
        for i, value in enumerate(FLAG_VALUES):
            header_row.cells[i + 1].text = value["label"].title()
        make_table_headers_bold(table)

        # Add category rows
        for row_idx, flag_name in enumerate(FLAG_NAMES):
            row = table.rows[row_idx + 1]
            row.cells[0].text = flag_name["label"]

            # Check each flag value for this category
            for col_idx, flag_value in enumerate(FLAG_VALUES):
                field_name = f"{flag_name['id']}_{flag_value['id']}"
                if field_name in data and data[field_name]:
                    row.cells[col_idx + 1].text = "Yes"
                else:
                    row.cells[col_idx + 1].text = "No"


def add_clinvar_settings_table(doc, data):
    """Add ClinVar settings as tables resembling the ClinvarPane.vue structure."""
    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    # Check if we have any ClinVar data
    overall_settings = ["require_in_clinvar", "clinvar_paranoid_mode"]
    interpretation_settings = [
        f'clinvar_include_{interp["id"]}' for interp in CLINVAR_INTERPRETATIONS
    ]

    has_overall = any(setting in data for setting in overall_settings)
    has_interpretations = any(setting in data for setting in interpretation_settings)

    if not has_overall and not has_interpretations:
        doc.add_paragraph("No ClinVar settings configured.")
        return

    # Add Overall Settings section
    if has_overall:
        doc.add_heading("Overall Settings", level=2)

        # Count how many overall settings we actually have
        present_overall_settings = [setting for setting in overall_settings if setting in data]

        table = doc.add_table(rows=len(present_overall_settings) + 1, cols=2)
        table.style = "Table Grid"

        # Header
        header_row = table.rows[0]
        header_row.cells[0].text = "Setting"
        header_row.cells[1].text = "Enabled"
        make_table_headers_bold(table)

        # Add all present overall settings (both True and False)
        row_idx = 1
        if "require_in_clinvar" in data:
            row = table.rows[row_idx]
            row.cells[0].text = "require ClinVar membership"
            row.cells[1].text = format_boolean_value(data["require_in_clinvar"])
            row_idx += 1

        if "clinvar_paranoid_mode" in data:
            row = table.rows[row_idx]
            row.cells[0].text = 'enable "paranoid" mode'
            row.cells[1].text = format_boolean_value(data["clinvar_paranoid_mode"])

    # Add ClinVar Interpretations section
    if has_interpretations:
        doc.add_heading("ClinVar Interpretations", level=2)

        # Create table for all interpretations (both enabled and disabled)
        present_interpretations = []
        for interp in CLINVAR_INTERPRETATIONS:
            field_name = f'clinvar_include_{interp["id"]}'
            if field_name in data:
                present_interpretations.append((interp, data[field_name]))

        if present_interpretations:
            table = doc.add_table(rows=len(present_interpretations) + 1, cols=2)
            table.style = "Table Grid"

            # Header
            header_row = table.rows[0]
            header_row.cells[0].text = "Interpretation"
            header_row.cells[1].text = "Included"
            make_table_headers_bold(table)

            # Add all present interpretations (both True and False)
            for row_idx, (interp, enabled) in enumerate(present_interpretations):
                row = table.rows[row_idx + 1]
                row.cells[0].text = f'{interp["label"]} variants'
                row.cells[1].text = format_boolean_value(enabled)


def _create_simple_prioritization_table(doc, title, setting_name, setting_value):
    """Create a simple two-row table for prioritization settings."""
    create_simple_table(doc, title, [(setting_name, format_boolean_value(setting_value))])


def _create_multi_row_prioritization_table(doc, title, rows_data):
    """Create a multi-row table for prioritization settings."""
    if not rows_data:
        return
    create_simple_table(doc, title, rows_data)


def _add_phenotype_prioritization_section(doc, data):
    """Add phenotype prioritization section."""
    phenotype_rows = []
    if "prio_enabled" in data:
        phenotype_rows.append(
            ("Enable phenotype-based prioritization", format_boolean_value(data["prio_enabled"]))
        )
    if "prio_algorithm" in data:
        algorithm_label = PHENOTYPE_ALGORITHM_OPTIONS.get(
            data["prio_algorithm"], data["prio_algorithm"]
        )
        phenotype_rows.append(("Phenotype Similarity Algorithm", algorithm_label))

    # Always include HPO terms row, even if empty or not present
    if "prio_hpo_terms" in data:
        hpo_terms = format_list_value(data["prio_hpo_terms"])
    else:
        hpo_terms = "(none)"
    phenotype_rows.append(("HPO Terms", hpo_terms))

    _create_multi_row_prioritization_table(doc, "Phenotype Prioritization", phenotype_rows)


def _add_pathogenicity_prioritization_section(doc, data):
    """Add pathogenicity prioritization section."""
    pathogenicity_rows = []
    if "patho_enabled" in data:
        pathogenicity_rows.append(
            (
                "Enable variant pathogenicity-based prioritization",
                format_boolean_value(data["patho_enabled"]),
            )
        )
    if "patho_score" in data:
        score_label = PATHOGENICITY_SCORE_OPTIONS.get(data["patho_score"], data["patho_score"])
        pathogenicity_rows.append(("Pathogenicity Score", score_label))

    _create_multi_row_prioritization_table(doc, "Pathogenicity Prioritization", pathogenicity_rows)


def add_prioritization_settings_table(doc, data):
    """Add prioritization settings as tables resembling the PrioritizationPane.vue structure."""
    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    # Check if we have any prioritization data
    phenotype_fields = ["prio_enabled", "prio_algorithm", "prio_hpo_terms"]
    pathogenicity_fields = ["patho_enabled", "patho_score"]
    face_fields = ["gm_enabled"]
    combined_fields = ["pedia_enabled"]

    has_phenotype = any(field in data for field in phenotype_fields)
    has_pathogenicity = any(field in data for field in pathogenicity_fields)
    has_face = any(field in data for field in face_fields)
    has_combined = any(field in data for field in combined_fields)

    if not (has_phenotype or has_pathogenicity or has_face or has_combined):
        doc.add_paragraph("No prioritization settings configured.")
        return

    # Add sections using helper functions
    if has_phenotype:
        _add_phenotype_prioritization_section(doc, data)

    if has_pathogenicity:
        _add_pathogenicity_prioritization_section(doc, data)

    if has_face and "gm_enabled" in data:
        _create_simple_prioritization_table(
            doc,
            "Face Prioritization",
            "Enable GestaltMatcher-based prioritization",
            data["gm_enabled"],
        )

    if has_combined and "pedia_enabled" in data:
        _create_simple_prioritization_table(
            doc,
            "Combined Prioritization",
            "Enable PEDIA based prioritization",
            data["pedia_enabled"],
        )


def add_genes_regions_settings_table(doc, data):
    """Add genes and regions settings as a table, always showing both settings even if empty."""
    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    # Always create a table with both gene allowlist and genomic regions
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"

    # Header
    header_row = table.rows[0]
    header_row.cells[0].text = "Setting"
    header_row.cells[1].text = "Value"
    make_table_headers_bold(table)

    # Gene Allow List row
    row1 = table.rows[1]
    row1.cells[0].text = "Gene Allow List"
    gene_allowlist = data.get("gene_allowlist", [])
    row1.cells[1].text = format_list_value(gene_allowlist, max_items=10)

    # Genomic Regions row
    row2 = table.rows[2]
    row2.cells[0].text = "Genomic Regions"
    genomic_regions = data.get("genomic_region", [])
    row2.cells[1].text = format_list_value(genomic_regions, max_items=5)


def add_other_settings_table(doc, data):
    """Add any remaining unprocessed settings as a clean table."""
    if not isinstance(data, dict):
        add_generic_section(doc, data)
        return

    if not data:
        doc.add_paragraph("No other settings found.")
        return

    # Create table for other settings
    table = doc.add_table(rows=len(data) + 1, cols=2)
    table.style = "Table Grid"

    # Header
    header_row = table.rows[0]
    header_row.cells[0].text = "Setting"
    header_row.cells[1].text = "Value"
    make_table_headers_bold(table)

    # Add each setting as a row
    for row_idx, (key, value) in enumerate(data.items()):
        row = table.rows[row_idx + 1]

        # Format setting name (convert underscore to spaces, title case)
        setting_name = format_display_name(key)
        row.cells[0].text = setting_name

        # Format value based on type
        if isinstance(value, list):
            row.cells[1].text = format_list_value(value)
        elif isinstance(value, bool):
            row.cells[1].text = format_boolean_value(value)
        elif isinstance(value, dict):
            row.cells[1].text = f"({len(value)} items)"
        else:
            row.cells[1].text = str(value)


def add_generic_section(doc, data):
    """Add a generic section for data that doesn't have special formatting."""
    if isinstance(data, dict):
        for key, value in data.items():
            if value is not None and value != "" and value != []:
                # Add subsection
                doc.add_heading(format_display_name(key), level=2)

                if isinstance(value, list) and value:
                    # Create bullet points for lists
                    for item in value:
                        p = doc.add_paragraph()
                        p.style = "List Bullet"
                        if isinstance(item, dict):
                            p.add_run(json.dumps(item, indent=2))
                        else:
                            p.add_run(str(item))
                elif isinstance(value, dict):
                    # Recursively handle nested dictionaries
                    for subkey, subvalue in value.items():
                        if subvalue is not None and subvalue != "" and subvalue != []:
                            p = doc.add_paragraph()
                            p.add_run(f"{format_display_name(subkey)}: ")
                            if isinstance(subvalue, list) and subvalue:
                                p.add_run("\n")
                                for item in subvalue:
                                    bullet_p = doc.add_paragraph()
                                    bullet_p.style = "List Bullet"
                                    bullet_p.add_run(str(item))
                            else:
                                p.add_run(str(subvalue))
                else:
                    # Simple value
                    p = doc.add_paragraph()
                    p.add_run(str(value))
    elif isinstance(data, list) and data:
        for item in data:
            p = doc.add_paragraph()
            p.style = "List Bullet"
            if isinstance(item, dict):
                p.add_run(json.dumps(item, indent=2))
            else:
                p.add_run(str(item))
    else:
        p = doc.add_paragraph()
        p.add_run(str(data))


def _setup_document_styles(doc):
    """
    Configure document styles for consistent formatting.

    Args:
        doc: The docx Document object to configure.
    """
    normal_style = doc.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Calibri"
    normal_font.size = Pt(11)


def _add_document_header(doc, case_info, request, filter_settings):
    """
    Add document title and metadata section.

    Args:
        doc: The docx Document object.
        case_info (dict): Information about the case.
        request: The HTTP request object.
        filter_settings (dict): The filter settings data.
    """
    title = doc.add_heading("Applied Filter Settings", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add metadata
    info_p = doc.add_paragraph()
    info_p.add_run("Export Date: ").bold = True
    info_p.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    # Add case name if available
    if case_info and case_info.get("name"):
        info_p.add_run("\nCase: ").bold = True
        info_p.add_run(case_info["name"])

    # Add user information
    if request.user and hasattr(request.user, "username"):
        info_p.add_run("\nUser: ").bold = True
        info_p.add_run(request.user.username)

    # Add transcript database information if available
    if "database" in filter_settings:
        db_value = filter_settings["database"]
        if db_value == "refseq":
            display_value = "RefSeq"
        elif db_value == "ensembl":
            display_value = "EnsEMBL"
        else:
            display_value = str(db_value)

        info_p.add_run("\nTranscript Database: ").bold = True
        info_p.add_run(display_value)


def _categorize_filter_settings(filter_settings):
    """
    Categorize filter settings into logical groups for organized document sections.

    Args:
        filter_settings (dict): The raw filter settings data.

    Returns:
        dict: Categorized settings organized by type (frequency, consequence, flags, etc.).
    """
    # Categorize settings
    frequency_settings = {}
    consequence_settings = {}
    flag_settings = {}
    clinvar_settings = {}
    prioritization_settings = {}
    genes_regions_settings = {}
    non_frequency_settings = {}

    for key, value in filter_settings.items():
        is_frequency = any(key.startswith(prefix) for prefix in FREQUENCY_PREFIXES)
        is_consequence = key in CONSEQUENCE_FIELDS
        is_flag = key in FLAG_FIELDS
        is_clinvar = key in CLINVAR_FIELDS
        is_prioritization = key in PRIORITIZATION_FIELDS
        is_genes_regions = key in GENES_REGIONS_FIELDS

        # Skip 'database' field as it's now handled in metadata
        if key == "database":
            continue
        elif is_frequency:
            frequency_settings[key] = value
        elif is_consequence:
            consequence_settings[key] = value
        elif is_flag:
            flag_settings[key] = value
        elif is_clinvar:
            clinvar_settings[key] = value
        elif is_prioritization:
            prioritization_settings[key] = value
        elif is_genes_regions:
            genes_regions_settings[key] = value
        else:
            non_frequency_settings[key] = value

    return {
        "frequency": frequency_settings,
        "consequence": consequence_settings,
        "flag": flag_settings,
        "clinvar": clinvar_settings,
        "prioritization": prioritization_settings,
        "genes_regions": genes_regions_settings,
        "non_frequency": non_frequency_settings,
    }


def _add_categorized_sections(doc, settings_dict, case_info, filter_settings):
    """
    Add categorized settings sections to the document in a logical order.

    Args:
        doc: The docx Document object.
        settings_dict (dict): Categorized filter settings.
        case_info (dict): Information about the case.
        filter_settings (dict): The original filter settings data.
    """
    section_mapping = {
        "quality": "Quality Settings",
        "frequency": "Frequency Settings",
        "consequence": "Variants & Effects Settings",
        "locus": "Locus Settings",
        "phenotypeprio": "Phenotype Priority Settings",
        "variantprio": "Variant Priority Settings",
        "clinvar": "ClinVar Settings",
        "columns": "Column Settings",
        "flags": "Flag Settings",
        "genotype": "Genotype Settings",
        "inheritance": "Inheritance Settings",
        "genomic_region": "Genomic Region",
        "gene_allowlist": "Gene Allowlist",
        "gene_blocklist": "Gene Blocklist",
        "transcript_flags": "Transcript Flags",
        "svtypes": "Structural Variant Types",
        "effects": "Effects",
        "max_exon_dist": "Maximum Exon Distance",
        "prio_enabled": "Prioritization Enabled",
        "prio_algorithm": "Prioritization Algorithm",
        "prio_hpo_terms": "HPO Terms",
        "patho_enabled": "Pathogenicity Enabled",
        "patho_score": "Pathogenicity Score",
    }

    # Follow the order of navigation tabs in FilterForm.vue:
    # 1. Genotype tab
    if "genotype" in settings_dict["non_frequency"]:
        add_genotype_settings_table_with_pedigree(
            doc, "Genotype Settings", settings_dict["non_frequency"]["genotype"], case_info
        )

    # 2. Frequency tab
    if settings_dict["frequency"]:
        add_key_value_table(doc, "Frequency Settings", settings_dict["frequency"])

    # 3. Prioritization tab
    if settings_dict["prioritization"]:
        add_key_value_table(doc, "Prioritization Settings", settings_dict["prioritization"])

    # 4. Variants & Effects tab (Consequence)
    if settings_dict["consequence"]:
        add_key_value_table(doc, "Variants & Effects Settings", settings_dict["consequence"])

    # 5. Quality tab
    if "quality" in settings_dict["non_frequency"]:
        add_key_value_table(doc, "Quality Settings", settings_dict["non_frequency"]["quality"])

    # 6. ClinVar tab
    if settings_dict["clinvar"]:
        add_key_value_table(doc, "ClinVar Settings", settings_dict["clinvar"])

    # 7. Gene Lists & Regions tab
    if settings_dict["genes_regions"]:
        add_key_value_table(doc, "Gene Lists & Regions Settings", settings_dict["genes_regions"])

    # 8. Flags & Comments tab
    if settings_dict["flag"]:
        add_key_value_table(doc, "Flags & Comments Settings", settings_dict["flag"])

    # Add any remaining standard sections (locus, inheritance, etc.)
    remaining_standard_sections = ["locus", "inheritance"]
    for section_key in remaining_standard_sections:
        if section_key in settings_dict["non_frequency"]:
            section_title = section_mapping.get(section_key, section_key.replace("_", " ").title())
            add_key_value_table(doc, section_title, settings_dict["non_frequency"][section_key])

    # Add remaining sections that don't fit into the main navigation
    processed_keys = {"genotype", "quality", "locus", "inheritance"}
    for key, value in settings_dict["non_frequency"].items():
        if key not in processed_keys and value is not None and value != "" and value != []:
            section_title = section_mapping.get(key, key.replace("_", " ").title())
            add_key_value_table(doc, section_title, value)

    # Add any remaining unprocessed settings as "Other Settings"
    remaining_settings = {}
    processed_standard_keys = {"genotype", "quality", "locus", "inheritance"}

    for key, value in filter_settings.items():
        # Check if this key was already processed by any of the specialized sections
        if (
            key not in settings_dict["frequency"]
            and key not in settings_dict["consequence"]
            and key not in settings_dict["flag"]
            and key not in settings_dict["clinvar"]
            and key not in settings_dict["prioritization"]
            and key != "database"  # Skip database field
            and key not in settings_dict["genes_regions"]
            and key not in processed_standard_keys
        ):
            # Only include if it has a meaningful value
            if value is not None and value != "" and value != []:
                remaining_settings[key] = value

    if remaining_settings:
        add_key_value_table(doc, "Other Settings", remaining_settings)


@csrf_exempt
@require_http_methods(["POST"])
def export_filter_settings(request):
    """Export filter settings as PDF file (with DOCX fallback)."""
    logger.info("Export filter settings endpoint called")

    try:
        # Parse request body
        if not request.body:
            return JsonResponse({"error": "Empty request body"}, status=400)

        data = json.loads(request.body)
        filter_settings = data.get("filter_settings", {})
        case_info = data.get("case_info", None)

        # Check permissions - if case info is provided, check case-based permissions
        # If no case info, just check authentication
        if case_info:
            perm_error = _check_case_permission(request, case_info)
            if perm_error:
                return perm_error
        else:
            auth_error = _check_user_authentication(request)
            if auth_error:
                return auth_error
        source = data.get("source", "unknown")
        export_format = data.get(
            "format", "docx"
        ).lower()  # Default to DOCX for backward compatibility

        logger.info(
            f"Processing filter settings export for source: {source}, format: {export_format}"
        )
        logger.debug(
            f"Filter settings keys: {list(filter_settings.keys()) if filter_settings else 'None'}"
        )
        logger.debug(f"Case info available: {case_info is not None}")

        # Validate format parameter
        if export_format not in ["pdf", "docx"]:
            return JsonResponse(
                {"error": f"Invalid format: {export_format}. Must be 'pdf' or 'docx'."}, status=400
            )

        # Create a new Document
        doc = Document()

        # Configure document styles
        _setup_document_styles(doc)

        # Add title and metadata
        _add_document_header(doc, case_info, request, filter_settings)

        # If no filter settings, add a note
        if not filter_settings:
            doc.add_paragraph("No filter settings data available to export.")
        else:
            # Categorize filter settings
            settings_dict = _categorize_filter_settings(filter_settings)

            # Add categorized sections to document
            _add_categorized_sections(doc, settings_dict, case_info, filter_settings)

        # Generate filename with case name if available
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d")
        base_filename = f"applied-filter-settings-{timestamp}"

        if case_info and case_info.get("name"):
            # Sanitize case name for filename
            case_name = str(case_info["name"]).strip()
            case_name = re.sub(r'[<>:"/\\|?*]', "_", case_name)
            case_name = re.sub(r"[^\w\-_\.]", "_", case_name)
            case_name = re.sub(r"_+", "_", case_name).strip("_")

            if case_name:  # Use case name if it's not empty after sanitization
                base_filename = f"applied-filter-settings-{case_name}-{timestamp}"

        # Generate file based on requested format
        if export_format == "pdf":
            # Try to generate PDF directly
            success, pdf_content, error_msg = generate_pdf_directly(
                filter_settings, case_info, request, base_filename
            )

            if success and pdf_content:
                # Return PDF response
                response = HttpResponse(
                    pdf_content,
                    content_type="application/pdf",
                )

                pdf_filename = f"{base_filename}.pdf"
                response["Content-Disposition"] = f'attachment; filename="{pdf_filename}"'

                logger.info(f"Successfully generated PDF file: {pdf_filename}")
                return response
            else:
                # Return error if PDF generation fails and was specifically requested
                logger.error(f"PDF generation failed: {error_msg}")
                return JsonResponse({"error": f"PDF generation failed: {error_msg}"}, status=500)
        else:
            # Generate DOCX
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)

            response = HttpResponse(
                bio.getvalue(),
                content_type=(
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
            )

            docx_filename = f"{base_filename}.docx"
            response["Content-Disposition"] = f'attachment; filename="{docx_filename}"'

            logger.info(f"Successfully generated DOCX file: {docx_filename}")
            return response

    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in request: {str(e)}")
        return JsonResponse({"error": "The request contained invalid JSON."}, status=400)
    except Exception as e:
        logger.error(f"Failed to export filter settings: {str(e)}", exc_info=True)
        return JsonResponse(
            {"error": "Internal server error while exporting filter settings."}, status=500
        )


def _clean_preset_dict(preset_dict):
    """Clean preset dictionary by removing metadata and unwanted fields."""
    if preset_dict is None:
        return {}

    cleaned = preset_dict.copy()

    # Remove metadata keys
    for key in ["sodar_uuid", "presetset", "date_created", "date_modified"]:
        cleaned.pop(key, None)

    # Remove frequency keys starting with specific prefixes
    keys_to_remove = [k for k in cleaned.keys() if k.startswith(("gene_blocklist",))]
    for key in keys_to_remove:
        cleaned.pop(key, None)

    return cleaned


def _add_preset_key_value_table(doc, data_dict):
    """Add a key-value table to the document for preset data."""
    if not data_dict:
        doc.add_paragraph("No additional configuration")
        return

    # Create table with 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Setting"
    hdr_cells[1].text = "Value"

    # Make header text bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add data rows
    for key, value in data_dict.items():
        if key != "label":  # Skip label as it's used as heading
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)

            # Handle different value types
            if isinstance(value, list):
                # Special handling for lists that should be bullet points
                if key in ["effects", "gene_allowlist", "genomic_region"]:
                    # Create bullet list in the cell
                    if value:  # Only if list is not empty
                        bullet_text = "\n".join([f"• {item}" for item in value])
                        row_cells[1].text = bullet_text
                    else:
                        row_cells[1].text = "(empty list)"
                else:
                    # For other lists, use JSON format
                    row_cells[1].text = json.dumps(value, indent=2)
            elif isinstance(value, dict):
                row_cells[1].text = json.dumps(value, indent=2)
            else:
                row_cells[1].text = str(value)


def _validate_preset_export_request(data):
    """Validate the request data for preset settings export."""
    project_uuid = data.get("project_uuid")
    presetset_uuid = data.get("presetset_uuid")

    if not project_uuid:
        return None, JsonResponse({"error": "project_uuid is required"}, status=400)
    if not presetset_uuid:
        return None, JsonResponse({"error": "presetset_uuid is required"}, status=400)

    return {
        "project_uuid": project_uuid,
        "presetset_uuid": presetset_uuid,
    }, None


def _find_presetset(project, presetset_uuid):
    """Find the presetset by UUID and return it with display label."""
    try:
        target_presetset = PresetSet.objects.get(sodar_uuid=presetset_uuid, project=project)
        return target_presetset, target_presetset.label, None
    except PresetSet.DoesNotExist:
        return None, None, "Preset set not found"


def _add_quickpreset_content(doc, quickpreset, preset_models):
    """Add content for a single quickpreset to the document."""
    # Add main heading with quickpreset label
    label = quickpreset.label or "Unnamed Preset"
    doc.add_heading(label, level=1)

    # Process inheritance (string value)
    inheritance = quickpreset.inheritance or ""
    if inheritance:
        doc.add_heading("Inheritance", level=2)
        doc.add_paragraph(inheritance)

    # Process other preset categories
    for model_class, category_key, category_title in preset_models:
        preset_obj = getattr(quickpreset, category_key, None)
        if preset_obj:
            # Add subheading with preset label
            preset_label = preset_obj.label or f"Unnamed {category_title}"
            doc.add_heading(f'{category_title}: "{preset_label}"', level=2)

            # Get preset data and clean it
            preset_dict = model_to_dict(preset_obj, exclude=["id", "date_created", "date_modified"])
            cleaned_data = _clean_preset_dict(preset_dict)
            _add_preset_key_value_table(doc, cleaned_data)

            # Add some spacing
            doc.add_paragraph()


def _create_preset_document(target_presetset, presetset_display_label):
    """Create the DOCX document with preset settings."""
    doc = Document()
    _setup_document_styles(doc)
    doc.add_heading(f"Query Presets - {presetset_display_label}", 0)

    # Get preset models and their categories
    preset_models = [
        (ChromosomePresets, "chromosome", "Chromosomes"),
        (QualityPresets, "quality", "Quality"),
        (ImpactPresets, "impact", "Impact"),
        (FrequencyPresets, "frequency", "Frequency"),
        (FlagsEtcPresets, "flagsetc", "Flags"),
    ]

    quickpresets = QuickPresets.objects.filter(presetset=target_presetset)
    if not quickpresets.exists():
        doc.add_paragraph("No quick presets found for this preset set.")
    else:
        for quickpreset in quickpresets:
            _add_quickpreset_content(doc, quickpreset, preset_models)

    return doc


def _create_export_response(
    doc, presetset_display_label, target_presetset=None, export_format="pdf"
):
    """Create the HTTP response with specified format."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d")

    # Sanitize the preset set label for filename
    safe_label = re.sub(r'[<>:"/\\|?*]', "_", presetset_display_label)
    safe_label = re.sub(r"[^\w\-_\.]", "_", safe_label)
    safe_label = re.sub(r"_+", "_", safe_label).strip("_")

    base_filename = f"preset-settings-{safe_label}-{timestamp}"

    # Generate file based on requested format
    if export_format == "pdf":
        # Try to generate PDF directly (if we have the presetset object)
        if target_presetset:
            success, pdf_content, error_msg = generate_preset_pdf_directly(
                target_presetset, presetset_display_label, base_filename
            )

            if success and pdf_content:
                # Return PDF response
                response = HttpResponse(
                    pdf_content,
                    content_type="application/pdf",
                )

                pdf_filename = f"{base_filename}.pdf"
                response["Content-Disposition"] = f'attachment; filename="{pdf_filename}"'

                return response, pdf_filename

            # Return error if PDF was specifically requested but failed
            logger.error(f"PDF generation failed for preset settings: {error_msg}")
            raise Exception(f"PDF generation failed: {error_msg}")
        else:
            raise Exception("PDF generation requires presetset object")

    # Generate DOCX
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    response = HttpResponse(
        bio.getvalue(),
        content_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    )

    docx_filename = f"{base_filename}.docx"
    response["Content-Disposition"] = f'attachment; filename="{docx_filename}"'

    return response, docx_filename


@csrf_exempt
@require_http_methods(["POST"])
def export_preset_settings(request):
    """Export preset settings as DOCX file."""
    logger.info("Export preset settings endpoint called")

    try:
        # Parse and validate request
        if not request.body:
            return JsonResponse({"error": "Empty request body"}, status=400)

        data = json.loads(request.body)
        validated_data, error = _validate_preset_export_request(data)

        if error:
            return error

        # Check project permissions
        has_permission, project, perm_error = _check_project_permission_by_uuid(
            request, validated_data["project_uuid"]
        )
        if not has_permission:
            return perm_error

        # Extract format parameter
        export_format = data.get("format", "pdf").lower()  # Default to PDF

        # Validate format parameter
        if export_format not in ["pdf", "docx"]:
            return JsonResponse(
                {"error": f"Invalid format: {export_format}. Must be 'pdf' or 'docx'."}, status=400
            )

        # At this point validated_data is guaranteed to be a dict, not None
        assert validated_data is not None

        logger.info(
            f"Processing preset settings export for project: {validated_data['project_uuid']}, "
            f"presetset_uuid: {validated_data['presetset_uuid']}, format: {export_format}"
        )

        # Get the project
        try:
            src_project = Project.objects.get(sodar_uuid=validated_data["project_uuid"])
        except Project.DoesNotExist:
            return JsonResponse(
                {"error": f"Project not found: {validated_data['project_uuid']}"}, status=404
            )

        # Find the presetset
        target_presetset, presetset_display_label, error = _find_presetset(
            src_project, validated_data["presetset_uuid"]
        )
        if error:
            return JsonResponse({"error": error}, status=404)

        # Create document with preset data
        doc = _create_preset_document(target_presetset, presetset_display_label)

        # Create and return response with specified format
        response, filename = _create_export_response(
            doc, presetset_display_label, target_presetset, export_format
        )
        logger.info(f"Successfully generated file: {filename}")
        return response

    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in request: {str(e)}")
        return JsonResponse({"error": "The request contained invalid JSON."}, status=400)
    except Exception as e:
        logger.error(f"Failed to export preset settings: {str(e)}", exc_info=True)
        return JsonResponse(
            {"error": "Internal server error while exporting preset settings."}, status=500
        )
