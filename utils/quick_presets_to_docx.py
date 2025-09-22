import json

try:
    import click
except ImportError:
    print("Error: click library is required.")
    print("Install it with: pip install click")
    exit(1)

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Error: python-docx library is required.")
    print("Install it with: pip install python-docx")
    exit(1)

def find_presetset_uuid(data, presetset_label):
    """Find the UUID for a presetset based on its label."""
    presetsets = data.get("presetsets", [])
    for presetset in presetsets:
        if presetset.get("label") == presetset_label:
            return presetset.get("sodar_uuid")
    return None

@click.command()
@click.option(
    '--json-file', '-j',
    type=click.Path(exists=True, readable=True),
    required=True,
    help='Path to the JSON input file'
)
@click.option(
    '--presetset-label', '-p',
    type=str,
    required=True,
    help='Label of the presetset to generate document for'
)
@click.option(
    '--output', '-o',
    type=str,
    default='preset_configuration.docx',
    help='Output filename for the Word document'
)
def main(json_file, presetset_label, output):
    """Generate a Word document from presets JSON file for a specific presetset."""
    
    # Load the JSON file
    try:
        with open(json_file, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        click.echo(f"Error: JSON file '{json_file}' not found.", err=True)
        exit(1)
    except json.JSONDecodeError as e:
        click.echo(f"Error: Invalid JSON in file '{json_file}': {e}", err=True)
        exit(1)
    
    # Find the UUID for the specified presetset label
    target_uuid = find_presetset_uuid(data, presetset_label)
    if not target_uuid:
        available_labels = [ps.get("label", "Unnamed") for ps in data.get("presetsets", [])]
        click.echo(f"Error: Presetset with label '{presetset_label}' not found.", err=True)
        click.echo(f"Available presetset labels: {', '.join(available_labels)}", err=True)
        exit(1)
    
    click.echo(f"Found presetset '{presetset_label}' with UUID: {target_uuid}")
    
    # Create a new Word document
    doc = Document()

    # Customize document styles - set normal style font to match headings
    # Get the normal style and heading styles
    normal_style = doc.styles['Normal']

    # Set normal style font to match heading font (typically Calibri or similar)
    normal_font = normal_style.font

    # Copy font properties from heading to normal style
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)  # Set a readable size for body text

    doc.add_heading(f'Quick Presets - {presetset_label}', 0)

    # Helper function to clean preset dictionary (remove metadata)
    def clean_preset(preset_dict):
        if preset_dict is None:
            return {}
        
        cleaned = preset_dict.copy()
        
        # Remove metadata keys
        for key in ["sodar_uuid", "presetset", "date_created", "date_modified"]:
            cleaned.pop(key, None)
        
        # Remove frequency keys starting with specific prefixes
        keys_to_remove = [k for k in cleaned.keys() if k.startswith(("exac_", "thousand_genomes_", "gene_blocklist"))]
        for key in keys_to_remove:
            cleaned.pop(key, None)
        
        return cleaned

    # Helper function to add a key-value table to the document
    def add_key_value_table(doc, data_dict):
        if not data_dict:
            doc.add_paragraph("No additional configuration")
            return
        
        # Create table with 2 columns
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Setting'
        hdr_cells[1].text = 'Value'
        
        # Make header text bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for key, value in data_dict.items():
            if key != 'label':  # Skip label as it's used as heading
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                
                # Handle different value types
                if isinstance(value, list):
                    # Special handling for lists that should be bullet points
                    if key in ['effects', 'gene_allowlist', 'genomic_region']:
                        # Create bullet list in the cell
                        if value:  # Only if list is not empty
                            bullet_text = '\n'.join([f'• {item}' for item in value])
                            row_cells[1].text = bullet_text
                        else:
                            row_cells[1].text = "(empty list)"
                    else:
                        # For other lists, use JSON format
                        row_cells[1].text = json.dumps(value, indent=2)
                elif isinstance(value, dict):
                    row_cells[1].text = json.dumps(value, indent=2)
                else:
                    row_cells[1].text = str(value)

    # Get quickpresets and presets data
    quickpresets = data.get("quickpresets", {})
    presets = data.get("presets", {})

    # Process each quickpreset
    for quickpreset in quickpresets:
        # Filter by specific presetset if needed
        if not quickpreset.get("presetset") == target_uuid:
            continue

        # Add main heading with quickpreset label
        label = quickpreset.get("label", "Unnamed Preset")
        doc.add_heading(label, level=1)
        
        # Process inheritance (string value)
        inheritance = quickpreset.get("inheritance", "")
        if inheritance:
            doc.add_heading("Inheritance", level=2)
            doc.add_paragraph(inheritance)
        
        # Process other preset categories (dictionaries with labels)
        preset_categories = [
            ("frequency", "Frequency"),
            ("impact", "Impact"), 
            ("quality", "Quality"),
            ("chromosome", "Chromosomes"),
            ("flagsetc", "Flags")
        ]
        
        for category_key, category_title in preset_categories:
            preset_uuid = quickpreset.get(category_key, "")
            if not preset_uuid:
                continue
                
            # Find the preset by UUID
            preset_data = None
            category_presets = presets.get(category_key, [])
            
            # Handle both list and dict formats
            if isinstance(category_presets, list):
                for entry in category_presets:
                    if entry.get("sodar_uuid", "") == preset_uuid:
                        preset_data = entry
                        break
            elif isinstance(category_presets, dict):
                for entry in category_presets.values():
                    if entry.get("sodar_uuid", "") == preset_uuid:
                        preset_data = entry
                        break
            
            if preset_data:
                # Add subheading with preset label
                preset_label = preset_data.get("label", f"Unnamed {category_title}")
                doc.add_heading(f"{category_title}: \"{preset_label}\"", level=2)
                
                # Clean the preset data and add as table
                cleaned_data = clean_preset(preset_data)
                add_key_value_table(doc, cleaned_data)
                
                # Add some spacing
                doc.add_paragraph()

    # Save the document
    doc.save(output)
    click.echo(f"Document saved as {output}")


if __name__ == "__main__":
    main()
